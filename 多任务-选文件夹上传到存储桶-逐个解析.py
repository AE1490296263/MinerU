import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import requests
import time
import os
import threading
from pathlib import Path
import shutil
import tempfile
import uuid
import zipfile
import json
import re

# 引入百度云SDK
from baidubce.auth import bce_credentials
from baidubce import bce_client_configuration
from baidubce.services.bos.bos_client import BosClient

# =========================
#  百度 BOS 配置信息 (保留您的配置)
# =========================
BOS_ACCESS_KEY = "ALTAKdxOEonUtkoh3kpFIiXmup"
BOS_SECRET_KEY = "abd78056d7284da59d92528a0692645e"
BOS_BUCKET = "ae86-minerutest1"
BOS_ENDPOINT = "https://bj.bcebos.com" 

class BatchBOSConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("MinerU 批量转换工具 (BOS托管 + 批量API)")
        self.root.geometry("950x800")

        # MinerU API 信息
        self.token = "eyJ0eXBlIjoiSldUIiwiYWxnIjoiSFM1MTIifQ.eyJqdGkiOiI1MzAwODI3OSIsInJvbCI6IlJPTEVfUkVHSVNURVIiLCJpc3MiOiJPcGVuWExhYiIsImlhdCI6MTc2MzU0NDM5NSwiY2xpZW50SWQiOiJsa3pkeDU3bnZ5MjJqa3BxOXgydyIsInBob25lIjoiMTg0NjAzMDAxOTciLCJvcGVuSWQiOm51bGwsInV1aWQiOiI5NjY3ODRiNC0wMjRjLTQ3NzUtYjE5Ny1kZWY5NTIyZmJjZDciLCJlbWFpbCI6IiIsImV4cCI6MTc2NDc1Mzk5NX0.HPAoPC83v5Xi-ZxjTshshZljtR7zTyTyKAVSt4qSCfCCShaVKWE7_K1bC2lWNrZJWi8r-hpTbv8ym6uRKBCizg"
        
        self.urls = {
            "batch_task": "https://mineru.net/api/v4/extract/task/batch",
            "batch_query": "https://mineru.net/api/v4/extract-results/batch/{}"
        }

        self.output_dir = r"D:\Desktop\项目\MinerU输出\批量处理结果"
        
        # 状态追踪
        self.files_data = [] # [{'path': str, 'name': str, 'bos_url': str, 'data_id': str, 'status': str}]
        self.processed_ids = set()
        self.is_running = False

        self.setup_ui()
        self.init_bos_client()
        self.cleanup_old_temp_files()

    def setup_ui(self):
        main = ttk.Frame(self.root, padding="15")
        main.pack(fill=tk.BOTH, expand=True)

        # 1. 文件夹选择
        input_frame = ttk.LabelFrame(main, text="1. 选择包含PDF的文件夹", padding="10")
        input_frame.pack(fill=tk.X, pady=5)
        
        self.folder_path = tk.StringVar()
        ttk.Entry(input_frame, textvariable=self.folder_path).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Button(input_frame, text="浏览文件夹", command=self.browse_folder).pack(side=tk.LEFT)

        # 2. 转换选项
        opt_frame = ttk.LabelFrame(main, text="2. 转换配置", padding="10")
        opt_frame.pack(fill=tk.X, pady=5)
        
        self.model_version = tk.StringVar(value="vlm")
        ttk.Radiobutton(opt_frame, text="VLM模型", variable=self.model_version, value="vlm").pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(opt_frame, text="Layout模型", variable=self.model_version, value="layout").pack(side=tk.LEFT, padx=10)
        
        ttk.Separator(opt_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        
        self.enable_ocr = tk.BooleanVar(value=True)
        self.enable_formula = tk.BooleanVar(value=True)
        self.convert_to_word = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(opt_frame, text="OCR识别", variable=self.enable_ocr).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(opt_frame, text="公式识别", variable=self.enable_formula).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(opt_frame, text="转Word文档", variable=self.convert_to_word).pack(side=tk.LEFT, padx=5)

        # 3. 文件列表 (带滚动条)
        list_frame = ttk.LabelFrame(main, text="文件列表与状态", padding="10")
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Treeview 比 Listbox 更好显示状态
        columns = ("filename", "status", "info")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", selectmode="none")
        self.tree.heading("filename", text="文件名")
        self.tree.heading("status", text="状态")
        self.tree.heading("info", text="详情/BOS链接")
        
        self.tree.column("filename", width=200)
        self.tree.column("status", width=100)
        self.tree.column("info", width=400)
        
        vsb = ttk.Scrollbar(list_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        # 4. 日志与操作
        log_frame = ttk.LabelFrame(main, text="执行日志", padding="10")
        log_frame.pack(fill=tk.X, pady=5)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=8, font=("Consolas", 9), state='disabled')
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        btn_frame = ttk.Frame(main)
        btn_frame.pack(pady=10)
        
        self.btn_start = ttk.Button(btn_frame, text="开始批量处理", command=self.start_process, state='disabled')
        self.btn_start.pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="停止/退出", command=self.cleanup_and_quit).pack(side=tk.LEFT, padx=10)

    def init_bos_client(self):
        try:
            config = bce_client_configuration.BceClientConfiguration(
                credentials=bce_credentials.BceCredentials(BOS_ACCESS_KEY, BOS_SECRET_KEY),
                endpoint=BOS_ENDPOINT
            )
            self.bos_client = BosClient(config)
            self.log("✅ 百度BOS客户端初始化成功")
        except Exception as e:
            self.log(f"❌ BOS初始化失败: {e}")
            messagebox.showerror("错误", f"BOS初始化失败: {e}")

    def browse_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.folder_path.set(folder)
            self.scan_files(folder)

    def scan_files(self, folder):
        # 清空列表
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.files_data = []
        
        # 扫描PDF
        count = 0
        for root, dirs, files in os.walk(folder):
            for f in files:
                if f.lower().endswith(".pdf"):
                    full_path = os.path.join(root, f)
                    # 生成唯一ID用于Treeview追踪
                    data_id = f"fid_{uuid.uuid4().hex[:8]}"
                    
                    file_info = {
                        "path": full_path,
                        "name": f,
                        "data_id": data_id,
                        "bos_url": None,
                        "tree_id": data_id # 使用 data_id 作为 tree item id
                    }
                    self.files_data.append(file_info)
                    self.tree.insert("", tk.END, iid=data_id, values=(f, "等待", ""))
                    count += 1
        
        self.log(f"共加载 {count} 个PDF文件")
        if count > 0:
            self.btn_start.config(state='normal')

    # ==========================
    #  主逻辑线程
    # ==========================
    def start_process(self):
        if not self.files_data: return
        self.is_running = True
        self.btn_start.config(state='disabled')
        self.processed_ids.clear()
        
        threading.Thread(target=self.worker_thread, daemon=True).start()

    def worker_thread(self):
        try:
            # 1. 批量上传到 BOS
            upload_success_list = []
            total = len(self.files_data)
            
            self.log(">>> 阶段 1/3: 上传文件到百度云 BOS...")
            
            for idx, item in enumerate(self.files_data):
                if not self.is_running: break
                
                self.update_tree_status(item['tree_id'], "上传中...", "")
                
                bos_url = self.upload_single_to_bos(item['path'])
                
                if bos_url:
                    item['bos_url'] = bos_url
                    self.update_tree_status(item['tree_id'], "已上传", "准备提交")
                    upload_success_list.append(item)
                else:
                    self.update_tree_status(item['tree_id'], "上传失败", "BOS错误")
            
            if not upload_success_list:
                self.log("❌ 没有文件上传成功，任务终止")
                self.reset_ui()
                return

            # 2. 提交批量任务给 MinerU
            self.log(f">>> 阶段 2/3: 提交批量解析任务 ({len(upload_success_list)} 个文件)...")
            batch_id = self.submit_batch_task(upload_success_list)
            
            if not batch_id:
                self.log("❌ 提交API失败，请检查Token或网络")
                self.reset_ui()
                return
            
            self.log(f"✅ 批量任务提交成功，Batch ID: {batch_id}")
            
            # 3. 轮询结果并下载
            self.log(">>> 阶段 3/3: 轮询结果并处理...")
            self.poll_and_process_results(batch_id, len(upload_success_list))

        except Exception as e:
            self.log(f"❌ 发生未捕获异常: {e}")
            self.reset_ui()

    # ==========================
    #  核心功能函数
    # ==========================
    def upload_single_to_bos(self, file_path):
        """上传单个文件到BOS"""
        try:
            file_name = os.path.basename(file_path)
            object_key = f"batch_upload/{uuid.uuid4().hex[:8]}_{file_name}"
            
            self.bos_client.put_object_from_file(BOS_BUCKET, object_key, file_path)
            
            # 生成公共链接 (假设Bucket是公开读，或者生成签名URL)
            # 如果Bucket是私有的，需要 generate_presigned_url
            # 这里按您原代码逻辑，假设是公共读
            public_url = f"https://{BOS_BUCKET}.bj.bcebos.com/{object_key}"
            return public_url
        except Exception as e:
            self.log(f"BOS Upload Error [{os.path.basename(file_path)}]: {e}")
            return None

    def submit_batch_task(self, file_items):
        """构造批量Payload并提交"""
        files_payload = []
        for item in file_items:
            files_payload.append({
                "url": item['bos_url'],
                "data_id": item['data_id'] # 关键：用这个ID把结果和文件对应起来
            })
        
        data = {
            "files": files_payload,
            "model_version": self.model_version.get(),
            "enable_ocr": self.enable_ocr.get(),
            "enable_formula": self.enable_formula.get()
        }
        headers = {"Authorization": f"Bearer {self.token}", "Content-Type": "application/json"}
        
        try:
            resp = requests.post(self.urls['batch_task'], json=data, headers=headers, timeout=30)
            res = resp.json()
            if res.get('code') == 0:
                return res['data']['batch_id']
            else:
                self.log(f"API Error: {res.get('msg')}")
                return None
        except Exception as e:
            self.log(f"Req Error: {e}")
            return None

    def poll_and_process_results(self, batch_id, expected_count):
        url = self.urls['batch_query'].format(batch_id)
        headers = {"Authorization": f"Bearer {self.token}"}
        
        completed_count = 0
        
        while self.is_running:
            try:
                resp = requests.get(url, headers=headers, timeout=30)
                res = resp.json()
                
                if res.get('code') != 0:
                    time.sleep(5)
                    continue
                
                results = res['data'].get('extract_result', [])
                
                current_round_updates = 0
                
                for item in results:
                    data_id = item.get('data_id')
                    state = item.get('state')
                    
                    # 找到对应的 tree item
                    if data_id in self.processed_ids:
                        continue
                        
                    # 更新状态显示
                    if state == 'running':
                        self.update_tree_status(data_id, "解析中", "MinerU处理中...")
                    elif state == 'pending':
                        self.update_tree_status(data_id, "排队中", "等待资源...")
                    
                    if state == 'done':
                        dl_url = item.get('full_zip_url')
                        self.update_tree_status(data_id, "下载中", "解析完成，下载...")
                        
                        # 查找原始文件名用于文件夹命名
                        original_item = next((x for x in self.files_data if x['data_id'] == data_id), None)
                        original_name = original_item['name'] if original_item else "unknown.pdf"
                        
                        if self.download_extract_convert(dl_url, original_name):
                            self.update_tree_status(data_id, "✅ 完成", "已保存")
                        else:
                            self.update_tree_status(data_id, "❌ 失败", "下载/转换出错")
                            
                        self.processed_ids.add(data_id)
                        completed_count += 1
                        current_round_updates += 1
                        
                    elif state == 'failed':
                        err = item.get('err_msg', 'unknown')
                        self.update_tree_status(data_id, "❌ 失败", err)
                        self.processed_ids.add(data_id)
                        completed_count += 1
                        current_round_updates += 1
                
                if completed_count >= expected_count:
                    self.log("✨ 所有文件处理完毕！")
                    messagebox.showinfo("完成", f"批量处理完成！\n结果已保存在: {self.output_dir}")
                    self.reset_ui()
                    break
                
                time.sleep(5) # 轮询间隔
                
            except Exception as e:
                self.log(f"Polling Error: {e}")
                time.sleep(5)

    def download_extract_convert(self, url, original_filename):
        temp_zip = None
        try:
            # 准备路径
            safe_name = Path(original_filename).stem
            final_folder = os.path.join(self.output_dir, safe_name)
            if os.path.exists(final_folder): shutil.rmtree(final_folder)
            os.makedirs(final_folder, exist_ok=True)
            
            # 下载
            temp_zip = os.path.join(tempfile.gettempdir(), f"mineru_{uuid.uuid4().hex}.zip")
            with requests.get(url, stream=True) as r:
                r.raise_for_status()
                with open(temp_zip, 'wb') as f:
                    for chunk in r.iter_content(8192):
                        f.write(chunk)
            
            # 解压
            with zipfile.ZipFile(temp_zip, 'r') as z:
                z.extractall(final_folder)
            
            # 清理ZIP
            try: os.remove(temp_zip)
            except: pass
            
            # 转Word
            if self.convert_to_word.get():
                self.convert_md_to_word(final_folder)
            
            return True
        except Exception as e:
            self.log(f"Post-process error: {e}")
            return False

    # ==========================
    #  MD 转 Word (复用您的逻辑)
    # ==========================
    def convert_md_to_word(self, folder_path):
        try:
            # 查找MD
            md_files = []
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    if file.endswith('.md'):
                        md_files.append(os.path.join(root, file))
            
            if not md_files: return
            
            from docx import Document
            from docx.shared import Pt
            
            for md_file in md_files:
                word_file = md_file.replace('.md', '.docx')
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                content = self.process_latex_formulas(content)
                
                doc = Document()
                style = doc.styles['Normal']
                style.font.name = '宋体'
                style.font.size = Pt(10.5)
                
                for line in content.split('\n'):
                    line = line.strip()
                    if not line: continue
                    
                    if line.startswith('# '): doc.add_heading(line[2:], 1)
                    elif line.startswith('## '): doc.add_heading(line[3:], 2)
                    elif line.startswith('### '): doc.add_heading(line[4:], 3)
                    else:
                        p = doc.add_paragraph()
                        parts = self.split_text_and_formulas(line)
                        for part in parts:
                            if part.startswith('FORMULA:'):
                                run = p.add_run(part[8:])
                                run.italic = True
                            else:
                                p.add_run(part)
                
                doc.save(word_file)
                self.log(f"  Word转换生成: {os.path.basename(word_file)}")
        except Exception as e:
            self.log(f"Word转换失败: {e}")

    def process_latex_formulas(self, text):
        # 简单处理，复用您的逻辑
        text = re.sub(r'\$(.*?)\$', lambda m: f'FORMULA:{self.clean_latex(m.group(1))}', text)
        text = re.sub(r'\$\$(.*?)\$\$', lambda m: f'\nFORMULA:{self.clean_latex(m.group(1))}\n', text)
        return text

    def clean_latex(self, f):
        f = re.sub(r'\\mathrm\{(.*?)\}', r'\1', f)
        f = f.replace('\\sim', '~').replace('\\approx', '≈').replace('\\leq', '≤')
        f = f.replace('\\times', '×').replace('\\div', '÷').replace('\\', '')
        return f.strip()

    def split_text_and_formulas(self, text):
        parts = []
        curr = 0
        for m in re.finditer(r'FORMULA:(.*?)(?=FORMULA:|$)', text):
            if m.start() > curr: parts.append(text[curr:m.start()])
            parts.append(f"FORMULA:{m.group(1)}")
            curr = m.end()
        if curr < len(text): parts.append(text[curr:])
        return parts if parts else [text]

    # ==========================
    #  UI 辅助
    # ==========================
    def log(self, msg):
        t = time.strftime("%H:%M:%S")
        self.root.after(0, lambda: self._log_insert(f"[{t}] {msg}\n"))

    def _log_insert(self, msg):
        self.log_text.config(state='normal')
        self.log_text.insert(tk.END, msg)
        self.log_text.see(tk.END)
        self.log_text.config(state='disabled')

    def update_tree_status(self, item_id, status, info):
        self.root.after(0, lambda: self.tree.item(item_id, values=(self.tree.item(item_id)['values'][0], status, info)))

    def reset_ui(self):
        self.is_running = False
        self.btn_start.config(state='normal')

    def cleanup_old_temp_files(self):
        d = tempfile.gettempdir()
        for f in os.listdir(d):
            if f.startswith("mineru_"):
                try: os.remove(os.path.join(d, f))
                except: pass

    def cleanup_and_quit(self):
        self.is_running = False
        self.cleanup_old_temp_files()
        self.root.quit()
        self.root.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = BatchBOSConverter(root)
    root.protocol("WM_DELETE_WINDOW", app.cleanup_and_quit)
    root.mainloop()
