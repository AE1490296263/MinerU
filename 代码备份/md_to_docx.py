# md_to_docx.py
import os
import sys
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag

# ========== 在这里修改输入输出路径 ==========
# 示例（请用你自己的路径替换下面的值）：
MARKDOWN_PATH = r"D:\Desktop\项目\MinerU\MinerU输出\Emergency Medicine ( PDFDrive )_原始文件\full.md"
IMAGES_DIR = r"D:\Desktop\项目\MinerU\MinerU输出\Emergency Medicine ( PDFDrive )_原始文件\images"
OUTPUT_DOCX = r"D:\Desktop\项目\MinerU\MinerU输出\Emergency Medicine ( PDFDrive )_原始文件\output.docx"
# 最大图片宽度（英寸）
MAX_IMAGE_WIDTH_IN = 6.0
# ===========================================

def find_image_path(src, images_dir):
    if not src:
        return None
    src = src.strip()
    # 如果是绝对路径且存在，直接返回
    if os.path.isabs(src) and os.path.exists(src):
        return src
    # 直接在 images_dir 下拼接
    candidate = os.path.join(images_dir, src)
    if os.path.exists(candidate):
        return candidate
    # 有时候 src 只是文件名或包含 ../ 等，尝试 basename
    base = os.path.basename(src)
    candidate2 = os.path.join(images_dir, base)
    if os.path.exists(candidate2):
        return candidate2
    # 遍历 images_dir 递归匹配相同文件名
    for root, dirs, files in os.walk(images_dir):
        for f in files:
            if f == base or f == src:
                return os.path.join(root, f)
    # 未找到
    return None

def add_text_with_format(paragraph, text, bold=False, italic=False, underline=False, code=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def process_inline(element, paragraph, doc):
    # element can be NavigableString or Tag
    if isinstance(element, NavigableString):
        paragraph.add_run(str(element))
        return
    if isinstance(element, Tag):
        name = element.name.lower()
        if name in ('strong', 'b'):
            for child in element.children:
                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child)).bold = True
                else:
                    process_inline(child, paragraph, doc)
        elif name in ('em', 'i'):
            for child in element.children:
                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child)).italic = True
                else:
                    process_inline(child, paragraph, doc)
        elif name == 'code':
            # inline code
            run = paragraph.add_run(element.get_text())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif name == 'a':
            text = element.get_text()
            href = element.get('href', '')
            run = paragraph.add_run(text)
            # show url in parentheses if it's external or different
            if href and href != text:
                paragraph.add_run(f" ({href})")
        elif name == 'img':
            src = element.get('src') or element.get('data-src') or element.get('alt')
            img_path = find_image_path(src, IMAGES_DIR)
            if img_path and os.path.exists(img_path):
                run = paragraph.add_run()
                try:
                    run.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                except Exception:
                    # fallback: document-level add_picture then center paragraph
                    doc.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
            else:
                paragraph.add_run(f"[图片未找到: {src}]")
        else:
            # 默认递归所有子节点
            for child in element.children:
                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child))
                else:
                    process_inline(child, paragraph, doc)

def convert_html_to_docx(html, doc, images_dir):
    soup = BeautifulSoup(html, 'html.parser')
    body = soup
    for elem in body.children:
        if isinstance(elem, NavigableString):
            text = str(elem).strip()
            if text:
                doc.add_paragraph(text)
            continue
        if not isinstance(elem, Tag):
            continue
        tag = elem.name.lower()
        if tag in ('h1','h2','h3','h4','h5','h6'):
            level = int(tag[1])
            text = elem.get_text()
            try:
                doc.add_heading(text, level=level-1)
            except Exception:
                doc.add_heading(text, level=0)
        elif tag == 'p':
            p = doc.add_paragraph()
            imgs = elem.find_all('img', recursive=False)
            if len(imgs) == 1 and elem.get_text(strip=True) == '':
                img_tag = imgs[0]
                src = img_tag.get('src') or img_tag.get('data-src') or img_tag.get('alt')
                img_path = find_image_path(src, images_dir)
                if img_path and os.path.exists(img_path):
                    run = p.add_run()
                    try:
                        run.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception:
                        doc.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                else:
                    p.add_run(f"[图片未找到: {src}]")
            else:
                for child in elem.children:
                    process_inline(child, p, doc)
        elif tag in ('ul', 'ol'):
            is_ordered = (tag == 'ol')
            for li in elem.find_all('li', recursive=False):
                style = 'List Number' if is_ordered else 'List Bullet'
                p = doc.add_paragraph(style=style)
                for child in li.children:
                    process_inline(child, p, doc)
        elif tag in ('pre',):
            code_text = elem.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif tag == 'hr':
            doc.add_page_break()
        elif tag == 'img':
            p = doc.add_paragraph()
            src = elem.get('src') or elem.get('data-src') or elem.get('alt')
            img_path = find_image_path(src, images_dir)
            if img_path and os.path.exists(img_path):
                run = p.add_run()
                try:
                    run.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception:
                    doc.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
            else:
                p.add_run(f"[图片未找到: {src}]")
        elif tag == 'table':
            # 处理表格
            rows = elem.find_all('tr')
            if not rows:
                continue
            # 统计列数
            first_row = rows[0]
            cols = first_row.find_all(['td', 'th'])
            n_cols = len(cols)
            n_rows = len(rows)
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    cell_text = ''.join(cell.strings).strip()
                    table.cell(i, j).text = cell_text
        else:
            if elem.get_text(strip=True):
                p = doc.add_paragraph()
                for child in elem.children:
                    if isinstance(child, NavigableString):
                        p.add_run(str(child))
                    else:
                        process_inline(child, p, doc)
def main():
    if not os.path.exists(MARKDOWN_PATH):
        print(f"错误：Markdown 文件未找到：{MARKDOWN_PATH}")
        sys.exit(1)
    if not os.path.isdir(IMAGES_DIR):
        print(f"警告：图片文件夹未找到：{IMAGES_DIR} ; 脚本仍会尝试按路径查找图片。")

    with io.open(MARKDOWN_PATH, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 将 markdown 转为 HTML（支持代码块和表格）
    html = markdown.markdown(md_text, extensions=['fenced_code', 'codehilite', 'tables'])

    doc = Document()
    convert_html_to_docx(html, doc, IMAGES_DIR)

    # 保存
    out_dir = os.path.dirname(OUTPUT_DOCX)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"已生成 Word 文件：{OUTPUT_DOCX}")

if __name__ == '__main__':
    main()