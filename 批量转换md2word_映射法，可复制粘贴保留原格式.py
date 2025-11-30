# batch_md_to_docx.py
import os
import sys
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
import re

# ========== 在这里修改输入输出路径 ==========
# 示例（请用你自己的路径替换下面的值）：
INPUT_ROOT_DIR = r"D:\Desktop\项目\MinerU输出\题库文件夹"  # 根目录，包含所有年份子文件夹
# 最大图片宽度（英寸）
MAX_IMAGE_WIDTH_IN = 6.0
# ===========================================

def find_image_path(src, images_dir):
    if not src:
        return None
    src = src.strip()
    # 如果是绝对路径且存在，直接返回
    if os.path.isabs(src) and os.path.exists(src):
        return src
    # 直接在 images_dir 下拼接
    candidate = os.path.join(images_dir, src)
    if os.path.exists(candidate):
        return candidate
    # 有时候 src 只是文件名或包含 ../ 等，尝试 basename
    base = os.path.basename(src)
    candidate2 = os.path.join(images_dir, base)
    if os.path.exists(candidate2):
        return candidate2
    # 遍历 images_dir 递归匹配相同文件名
    for root, dirs, files in os.walk(images_dir):
        for f in files:
            if f == base or f == src:
                return os.path.join(root, f)
    # 未找到
    return None

def add_text_with_format(paragraph, text, bold=False, italic=False, underline=False, code=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def process_inline(element, paragraph, doc, images_dir):  # 添加 images_dir 参数
    # element can be NavigableString or Tag
    if isinstance(element, NavigableString):
        paragraph.add_run(str(element))
        return
    if isinstance(element, Tag):
        name = element.name.lower()
        if name in ('strong', 'b'):
            for child in element.children:
                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child)).bold = True
                else:
                    process_inline(child, paragraph, doc, images_dir)  # 传递 images_dir
        elif name in ('em', 'i'):
            for child in element.children:
                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child)).italic = True
                else:
                    process_inline(child, paragraph, doc, images_dir)  # 传递 images_dir
        elif name == 'code':
            # inline code
            run = paragraph.add_run(element.get_text())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif name == 'a':
            text = element.get_text()
            href = element.get('href', '')
            run = paragraph.add_run(text)
            # show url in parentheses if it's external or different
            if href and href != text:
                paragraph.add_run(f" ({href})")
        elif name == 'img':
            src = element.get('src') or element.get('data-src') or element.get('alt')
            img_path = find_image_path(src, images_dir)
            if img_path and os.path.exists(img_path):
                run = paragraph.add_run()
                try:
                    run.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                except Exception:
                    # fallback: document-level add_picture then center paragraph
                    doc.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
            else:
                paragraph.add_run(f"[图片未找到: {src}]")
        else:
            # 默认递归所有子节点
            for child in element.children:
                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child))
                else:
                    process_inline(child, paragraph, doc, images_dir)  # 传递 images_dir

def convert_html_to_docx(html, doc, images_dir):
    soup = BeautifulSoup(html, 'html.parser')
    body = soup
    for elem in body.children:
        if isinstance(elem, NavigableString):
            text = str(elem).strip()
            if text:
                doc.add_paragraph(text)
            continue
        if not isinstance(elem, Tag):
            continue
        tag = elem.name.lower()
        if tag in ('h1','h2','h3','h4','h5','h6'):
            level = int(tag[1])
            text = elem.get_text()
            try:
                doc.add_heading(text, level=level-1)
            except Exception:
                doc.add_heading(text, level=0)
        elif tag == 'p':
            p = doc.add_paragraph()
            imgs = elem.find_all('img', recursive=False)
            if len(imgs) == 1 and elem.get_text(strip=True) == '':
                img_tag = imgs[0]
                src = img_tag.get('src') or img_tag.get('data-src') or img_tag.get('alt')
                img_path = find_image_path(src, images_dir)
                if img_path and os.path.exists(img_path):
                    run = p.add_run()
                    try:
                        run.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception:
                        doc.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                else:
                    p.add_run(f"[图片未找到: {src}]")
            else:
                for child in elem.children:
                    process_inline(child, p, doc, images_dir)  # 传递 images_dir
        elif tag in ('ul', 'ol'):
            is_ordered = (tag == 'ol')
            for li in elem.find_all('li', recursive=False):
                style = 'List Number' if is_ordered else 'List Bullet'
                p = doc.add_paragraph(style=style)
                for child in li.children:
                    process_inline(child, p, doc, images_dir)  # 传递 images_dir
        elif tag in ('pre',):
            code_text = elem.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif tag == 'hr':
            doc.add_page_break()
        elif tag == 'img':
            p = doc.add_paragraph()
            src = elem.get('src') or elem.get('data-src') or elem.get('alt')
            img_path = find_image_path(src, images_dir)
            if img_path and os.path.exists(img_path):
                run = p.add_run()
                try:
                    run.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception:
                    doc.add_picture(img_path, width=Inches(MAX_IMAGE_WIDTH_IN))
            else:
                p.add_run(f"[图片未找到: {src}]")
        elif tag == 'table':
            # 处理表格
            rows = elem.find_all('tr')
            if not rows:
                continue
            # 统计列数
            first_row = rows[0]
            cols = first_row.find_all(['td', 'th'])
            n_cols = len(cols)
            n_rows = len(rows)
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    cell_text = ''.join(cell.strings).strip()
                    table.cell(i, j).text = cell_text
        else:
            if elem.get_text(strip=True):
                p = doc.add_paragraph()
                for child in elem.children:
                    if isinstance(child, NavigableString):
                        p.add_run(str(child))
                    else:
                        process_inline(child, p, doc, images_dir)  # 传递 images_dir

def preprocess_math(md_text):
    """
    将 Markdown 中的 LaTeX 风格数学表达式转换为可读文本
    """
    if not md_text:
        return md_text

    # 先把常见的转义百分号恢复
    md_text = md_text.replace(r'\%', '%')

    # 常见希腊字母和数学符号映射
    greek_letters = {
        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'Gamma': 'Γ',
        'delta': 'δ', 'Delta': 'Δ', 'epsilon': 'ε', 'varepsilon': 'ε',
        'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'Theta': 'Θ',
        'iota': 'ι', 'kappa': 'κ', 'lambda': 'λ', 'Lambda': 'Λ',
        'mu': 'μ', 'nu': 'ν', 'xi': 'ξ', 'Xi': 'Ξ',
        'pi': 'π', 'Pi': 'Π', 'rho': 'ρ', 'sigma': 'σ', 'Sigma': 'Σ',
        'tau': 'τ', 'upsilon': 'υ', 'phi': 'φ', 'Phi': 'Φ',
        'chi': 'χ', 'psi': 'ψ', 'Psi': 'Ψ', 'omega': 'ω', 'Omega': 'Ω'
    }
    
    # 常见数学符号映射
    math_symbols = {
        'times': '×', 'div': '÷', 'pm': '±', 'mp': '∓',
        'cdot': '·', 'ast': '*', 'star': '⋆',
        'leq': '≤', 'geq': '≥', 'neq': '≠', 'equiv': '≡',
        'approx': '≈', 'sim': '∼', 'simeq': '≃', 'cong': '≅',
        'propto': '∝', 'infty': '∞', 'forall': '∀', 'exists': '∃',
        'nabla': '∇', 'partial': '∂', 'emptyset': '∅', 'varnothing': '∅',
        'in': '∈', 'notin': '∉', 'subset': '⊂', 'subseteq': '⊆',
        'supset': '⊃', 'supseteq': '⊇', 'cup': '∪', 'cap': '∩',
        'bigcup': '⋃', 'bigcap': '⋂', 'oplus': '⊕', 'otimes': '⊗',
        'wedge': '∧', 'vee': '∨', 'neg': '¬', 'Rightarrow': '⇒',
        'rightarrow': '→', 'leftarrow': '←', 'Leftrightarrow': '⇔',
        'leftrightarrow': '↔', 'mapsto': '↦', 'to': '→',
        'circ': '∘', 'bullet': '•', 'dagger': '†', 'ddagger': '‡',
        'perp': '⊥', 'parallel': '∥', 'angle': '∠', 'triangle': '△',
        'square': '□', 'diamond': '◇', 'clubsuit': '♣', 'diamondsuit': '♦',
        'heartsuit': '♥', 'spadesuit': '♠',
        'prime': "'"  # 添加这个映射
    }
    
    # 组合所有符号映射
    symbol_map = {**greek_letters, **math_symbols}

    def replace_math_expression(math_content):
        """替换数学表达式中的符号"""
        content = math_content.strip()
        
        # 首先特殊处理：5^{\prime} -> 5'（优先处理这种情况）
        content = re.sub(r'(\d+)\^\\?\{?\\?prime\}?', r"\1'", content)
        
        # 处理希腊字母和数学符号
        for cmd, symbol in symbol_map.items():
            content = re.sub(r'\\' + cmd + r'(?![a-zA-Z])', symbol, content)
        
        # 处理其他上标和下标（但要排除已经处理过的数字+撇号情况）
        content = re.sub(r'\^\{([^}]*)\}', r'^\1', content)
        content = re.sub(r'_\{([^}]*)\}', r'_\1', content)
        content = re.sub(r'\^([^{])', r'^\1', content)
        content = re.sub(r'_([^{])', r'_\1', content)
        
        # 处理分数 \frac{a}{b} -> a/b
        content = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'\1/\2', content)
        
        # 处理根号 \sqrt{a} -> √a
        content = re.sub(r'\\sqrt\{([^}]*)\}', r'√\1', content)
        
        # 处理数学字体命令
        content = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', content)
        content = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', content)
        content = re.sub(r'\\mathit\{([^}]*)\}', r'\1', content)
        content = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', content)
        
        # 删除残留的大括号
        content = content.replace('{', '').replace('}', '')
        
        # 压缩空白
        content = re.sub(r'\s+', ' ', content).strip()
        
        return content

    # 处理显示数学 $$...$$
    md_text = re.sub(r'\$\$(.+?)\$\$', 
                    lambda m: replace_math_expression(m.group(1)), 
                    md_text, flags=re.S)
    
    # 处理行内数学 $...$
    md_text = re.sub(r'\$(.+?)\$', 
                    lambda m: replace_math_expression(m.group(1)), 
                    md_text, flags=re.S)
    
    return md_text


def convert_single_file(md_file_path, output_docx_path):
    """转换单个 markdown 文件"""
    try:
        # 获取对应的 images 文件夹路径
        folder_path = os.path.dirname(md_file_path)
        images_dir = os.path.join(folder_path, "images")
        
        with io.open(md_file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()

        # 预处理 LaTeX 风格的数学表达式与转义百分号
        md_text = preprocess_math(md_text)

        # 将 markdown 转为 HTML（支持代码块和表格）
        html = markdown.markdown(md_text, extensions=['fenced_code', 'codehilite', 'tables'])

        doc = Document()
        convert_html_to_docx(html, doc, images_dir)

        # 保存
        out_dir = os.path.dirname(output_docx_path)
        if out_dir and not os.path.exists(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        doc.save(output_docx_path)
        
        return True, f"成功转换: {output_docx_path}"
        
    except Exception as e:
        return False, f"转换失败 {md_file_path}: {str(e)}"

def find_md_files(root_dir):
    """查找所有子文件夹中的 full.md 文件"""
    md_files = []
    for root, dirs, files in os.walk(root_dir):
        if "full.md" in files:
            md_path = os.path.join(root, "full.md")
            # 生成对应的输出路径
            relative_path = os.path.relpath(root, root_dir)
            output_filename = f"{os.path.basename(root)}.docx"
            output_path = os.path.join(root, output_filename)
            md_files.append((md_path, output_path))
    return md_files

def main():
    if not os.path.exists(INPUT_ROOT_DIR):
        print(f"错误：根目录未找到：{INPUT_ROOT_DIR}")
        sys.exit(1)

    print(f"开始扫描目录: {INPUT_ROOT_DIR}")
    md_files = find_md_files(INPUT_ROOT_DIR)
    
    if not md_files:
        print("未找到任何 full.md 文件")
        return
    
    print(f"找到 {len(md_files)} 个 full.md 文件")
    
    success_count = 0
    failed_files = []
    
    for i, (md_path, output_path) in enumerate(md_files, 1):
        print(f"\n[{i}/{len(md_files)}] 正在转换: {md_path}")
        
        success, message = convert_single_file(md_path, output_path)
        
        if success:
            print(f"✓ {message}")
            success_count += 1
        else:
            print(f"✗ {message}")
            failed_files.append((md_path, message))
    
    # 输出总结
    print(f"\n{'='*50}")
    print(f"转换完成!")
    print(f"成功: {success_count}/{len(md_files)}")
    print(f"失败: {len(failed_files)}/{len(md_files)}")
    
    if failed_files:
        print("\n失败的文件:")
        for md_path, error in failed_files:
            print(f"  - {md_path}: {error}")

if __name__ == '__main__':
    main()
